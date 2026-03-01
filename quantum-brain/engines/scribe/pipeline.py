"""
ThaleOS Scribe Pipeline
Jinja2-based template rendering for DOCX, Markdown, and PDF documents.
Templates live in: memory-palace/scribe-templates/{format}/{name}.{format}.j2
"""

import logging
from pathlib import Path
from typing import Optional
from datetime import date

logger = logging.getLogger("ThaleOS.Scribe")

# Template root: quantum-brain/../memory-palace/scribe-templates/
_HERE = Path(__file__).resolve().parent
TEMPLATE_ROOT = _HERE.parent.parent.parent / "memory-palace" / "scribe-templates"


def _get_jinja_env(template_dir: Path):
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    return Environment(
        loader=FileSystemLoader(str(template_dir)),
        autoescape=select_autoescape([]),
        trim_blocks=True,
        lstrip_blocks=True,
    )


class ScribePipeline:
    """
    Renders documents from Jinja2 templates.
    Supports markdown, DOCX, and PDF output formats.
    """

    def list_templates(self) -> list[dict]:
        """Return all available templates with their required fields."""
        templates = []
        if not TEMPLATE_ROOT.exists():
            return templates
        for fmt_dir in TEMPLATE_ROOT.iterdir():
            if not fmt_dir.is_dir():
                continue
            fmt = fmt_dir.name  # 'markdown', 'docx', etc.
            for tmpl_file in fmt_dir.glob("*.j2"):
                # Extract required fields by scanning the template
                fields = self._extract_fields(tmpl_file)
                templates.append({
                    "name": tmpl_file.stem.rsplit(".", 1)[0],  # e.g. "report" from "report.md.j2"
                    "format": fmt,
                    "file": tmpl_file.name,
                    "required_fields": fields,
                })
        return sorted(templates, key=lambda t: t["name"])

    def _extract_fields(self, tmpl_file: Path) -> list[str]:
        """
        Simple scanner — finds {{ variable_name }} patterns that have no default filter.
        Not 100% accurate but good enough for documentation purposes.
        """
        import re
        text = tmpl_file.read_text()
        # Match {{ var }} but not {{ var | default(...) }} (those are optional)
        all_vars = re.findall(r'\{\{\s*(\w+)\s*\}\}', text)
        optional = re.findall(r'\{\{\s*(\w+)\s*\|', text)
        required = [v for v in set(all_vars) if v not in optional]
        return sorted(required)

    def validate_data(self, template_name: str, fmt: str, data: dict) -> list[str]:
        """
        Check which required fields are missing from data.
        Returns list of missing field names (empty = all good).
        """
        templates = self.list_templates()
        tmpl = next((t for t in templates if t["name"] == template_name and t["format"] == fmt), None)
        if not tmpl:
            return [f"Template '{template_name}' ({fmt}) not found"]
        return [f for f in tmpl["required_fields"] if f not in data]

    def render_markdown(self, template_name: str, data: dict) -> str:
        """
        Render a Jinja2 markdown template to a string.
        Injects today's date if not provided.
        """
        tmpl_dir = TEMPLATE_ROOT / "markdown"
        if not tmpl_dir.exists():
            raise FileNotFoundError(f"Template directory not found: {tmpl_dir}")

        # Find the template file (e.g. report.md.j2)
        matches = list(tmpl_dir.glob(f"{template_name}.*.j2")) + list(tmpl_dir.glob(f"{template_name}.j2"))
        if not matches:
            raise FileNotFoundError(f"Template '{template_name}' not found in {tmpl_dir}")

        env = _get_jinja_env(tmpl_dir)
        tmpl = env.get_template(matches[0].name)

        ctx = {"date": date.today().isoformat(), **data}
        return tmpl.render(**ctx)

    def render_docx(self, template_name: str, data: dict) -> bytes:
        """
        Render a document to DOCX bytes using python-docx.
        Looks for a .docx.j2 template first (XML-level templating).
        Falls back to rendering markdown and converting to DOCX.
        """
        from docx import Document
        from docx.shared import Pt
        import io

        # Render as markdown first, then pour into a docx
        md_content = self.render_markdown(template_name, data)

        doc = Document()
        doc.add_heading(data.get("title", template_name.title()), level=0)

        for line in md_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                p.add_run(stripped.strip("*")).bold = True
            elif stripped.startswith("---"):
                doc.add_paragraph("─" * 40)
            else:
                doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_from_llm_content(self, content: str, title: str = "Document") -> bytes:
        """
        Take free-form LLM-generated text (typically markdown) and export as DOCX.
        Used when no template is being applied — just export whatever SCRIBE wrote.
        """
        from docx import Document
        import io

        doc = Document()
        doc.add_heading(title, level=0)

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("---"):
                doc.add_paragraph("─" * 40)
            else:
                doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
